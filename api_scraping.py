# FastAPI
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Selenium e WebDriver
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager

# Manipulação de Documentos
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from fastapi.responses import FileResponse

# Utilitários Python
from datetime import datetime
import logging
import time
import asyncio
import os
from typing import List, Dict, Optional


# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuração da pasta de exports
EXPORTS_DIR = 'exports'
os.makedirs(EXPORTS_DIR, exist_ok=True)

# Configuração da API
app = FastAPI(title="Web Scraping API")

# Configuração CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ScrapingRequest(BaseModel):
    url: str
    seletores: dict
    nome_arquivo: str
    formato_arquivo: str
    limite_itens: int = 100
    limite_paginas: int = 5
    delay_entre_paginas: float = 2.0
    seletor_proxima_pagina: Optional[str] = None
    credenciais: Optional[Dict[str, str]] = None

# Em api_scraping.py

class ScrapingManager:
    def __init__(self):
        self.setup_chrome_options()

    def format_value(self, value):
        """Formata valores para exibição no documento"""
        if pd.isna(value):
            return ''
        elif isinstance(value, (int, float)):
            return f"{value:,}"
        else:
            return str(value)

    def setup_chrome_options(self):
        """Configura opções do Chrome"""
        options = Options()
        
        # Configurações headless
        options.add_argument('--headless=new')
        options.add_argument('--disable-gpu')
        
        # Configurações de performance
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-software-rasterizer')
        options.add_argument('--disable-extensions')
        options.add_argument('--window-size=1920,1080')
        
        # Configurações anti-detecção
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_experimental_option('excludeSwitches', ['enable-automation'])
        options.add_experimental_option('useAutomationExtension', False)
        
        # User agent
        options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36')
        
        # Preferências
        prefs = {
            'profile.default_content_setting_values.notifications': 2,
            'profile.managed_default_content_settings.images': 1,
            'disk-cache-size': 4096
        }
        options.add_experimental_option('prefs', prefs)
        
        self.chrome_options = options
        return options

    async def scrape_page(self, driver, url: str, seletores: dict) -> List[dict]:
        """Realiza o scraping de uma página específica"""
        results = []
        
        try:
            # Aguarda a página carregar completamente
            driver.execute_script("return document.readyState") == "complete"
            
            # Espera explícita pelo container
            try:
                container = WebDriverWait(driver, 15).until(
                    EC.presence_of_all_elements_located((By.CSS_SELECTOR, seletores['container']))
                )
                logger.info(f"Container encontrado com {len(container)} elementos")
                
                # Scroll até o final da página para garantir carregamento
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                await asyncio.sleep(2)  # Aguarda carregamento dinâmico
                
                # Encontra os elementos novamente após o scroll
                elementos = driver.find_elements(By.CSS_SELECTOR, seletores['container'])
                logger.info(f"Total de elementos após scroll: {len(elementos)}")
                
                for elemento in elementos:
                    item = {}
                    
                    # Extrai os dados de cada elemento
                    for campo, seletor in seletores.items():
                        if campo != 'container':
                            try:
                                # Espera pelo elemento específico
                                valor = WebDriverWait(elemento, 5).until(
                                    EC.presence_of_element_located((By.CSS_SELECTOR, seletor))
                                ).text.strip()
                                item[campo] = valor
                                logger.debug(f"Campo {campo} extraído: {valor}")
                            except Exception as e:
                                logger.warning(f"Erro ao extrair {campo}: {str(e)}")
                                item[campo] = ''
                    
                    if any(item.values()):  # Só adiciona se tiver algum valor
                        results.append(item)
                        logger.info(f"Item extraído: {item}")
                
            except TimeoutException:
                logger.error("Timeout ao aguardar elementos do container")
                raise
                
            logger.info(f"Total de resultados coletados na página: {len(results)}")
            return results
            
        except Exception as e:
            logger.error(f"Erro durante scraping da página: {str(e)}")
            raise

    async def login(self, driver, url, credenciais):
        """Realiza o login no sistema SED"""
        logger.info("Iniciando processo de login")
        driver.get(url)
        
        try:
            # Aguardar até que os campos de login estejam presentes
            username_field = WebDriverWait(driver, 90).until(
                EC.presence_of_element_located((By.ID, "Usuario"))
            )
            password_field = WebDriverWait(driver, 90).until(
                EC.presence_of_element_located((By.ID, "Senha"))
            )
            
            # Preencher campos de login
            username_field.send_keys(credenciais['usuario'])
            password_field.send_keys(credenciais['senha'])
            
            # Clicar no botão de login
            login_button = WebDriverWait(driver, 90).until(
                EC.element_to_be_clickable((By.ID, "btnAcessar"))
            )
            login_button.click()
            
            # Aguardar a página carregar após o login
            WebDriverWait(driver, 90).until(
                EC.url_contains("MinhasTurmas/GridAcesso")
            )
            logger.info("Login realizado com sucesso")
            
        except TimeoutException:
            logger.error("Erro ao fazer login: Elementos não encontrados ou página não carregou")
            raise HTTPException(
                status_code=500,
                detail={
                    "message": "Erro durante o login",
                    "error": "Elementos não encontrados ou página não carregou"
                }
            )

    async def process_all_pages(self, request: ScrapingRequest) -> Dict:
        """Processa múltiplas páginas de resultados"""
        all_results = []
        driver = None
        
        try:
            # Configuração inicial do WebDriver
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=self.chrome_options)
            driver.set_page_load_timeout(30)
            driver.implicitly_wait(10)
            
            # Verificar credenciais
            if request.credenciais:
                await self.login(driver, request.url, request.credenciais)
            
            # Inicializar variáveis
            current_url = request.url
            page = 1
            total_items = 0
            
            # Acessar URL inicial
            logger.info(f"Acessando URL: {current_url}")
            driver.get(current_url)
            await asyncio.sleep(5)  # Aguarda carregamento inicial
            
            # Loop de coleta de dados
            while (page <= request.limite_paginas and total_items < request.limite_itens):
                logger.info(f"Processando página {page}")
                
                try:
                    # Coletar dados da página atual
                    page_results = await self.scrape_page(driver, current_url, request.seletores)
                    if page_results:
                        all_results.extend(page_results)
                        total_items = len(all_results)
                        logger.info(f"Total acumulado: {total_items} itens")
                    
                    # Verificar próxima página
                    if request.seletor_proxima_pagina and total_items < request.limite_itens:
                        try:
                            next_button = WebDriverWait(driver, 10).until(
                                EC.element_to_be_clickable((By.CSS_SELECTOR, request.seletor_proxima_pagina))
                            )
                            driver.execute_script("arguments[0].scrollIntoView(true);", next_button)
                            await asyncio.sleep(1)
                            next_button.click()
                            await asyncio.sleep(request.delay_entre_paginas)
                            page += 1
                            current_url = driver.current_url
                        except Exception as e:
                            logger.info(f"Não há mais páginas: {str(e)}")
                            break
                    else:
                        break
                        
                except Exception as e:
                    logger.error(f"Erro ao processar página {page}: {str(e)}")
                    break

            if not all_results:
                raise HTTPException(
                    status_code=500,
                    detail="Nenhum dado foi coletado com os seletores fornecidos"
                )

            # Criar DataFrame
            df = pd.DataFrame(all_results)
            
            # Criar diretório se não existir
            os.makedirs(EXPORTS_DIR, exist_ok=True)
            
            try:
                file_name = request.nome_arquivo
                file_format = request.formato_arquivo.lower()
                
                file_path = os.path.join(EXPORTS_DIR, f"{file_name}.{file_format}")
                
                if file_format == 'xlsx':
                    # Salvar XLSX
                    df.to_excel(file_path, index=False, engine='openpyxl')
                    logger.info(f"Arquivo XLSX salvo: {file_path}")
                elif file_format == 'csv':
                    # Salvar CSV
                    df.to_csv(file_path, index=False, encoding='utf-8-sig')
                    logger.info(f"Arquivo CSV salvo: {file_path}")
                elif file_format == 'docx':
                    # Criar e salvar DOCX
                    doc = Document()
                    
                    # Adicionar título
                    title = doc.add_heading('Relatório de Web Scraping', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Adicionar informações gerais
                    doc.add_paragraph()
                    doc.add_paragraph(f'Data da extração: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
                    doc.add_paragraph(f'URL: {request.url}')
                    doc.add_paragraph(f'Total de itens extraídos: {len(all_results)}')
                    doc.add_paragraph(f'Total de páginas processadas: {page}')
                    doc.add_paragraph()

                    # Adicionar conteúdo do DataFrame
                    for _, row in df.iterrows():
                        # Primeiro adicionar o título (se existir)
                        if 'Título' in row:
                            titulo = str(row['Título']).strip()
                            if titulo:
                                heading = doc.add_heading(titulo, level=1)
                                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Depois adicionar os outros campos
                        for col in df.columns:
                            if col.lower() != 'título':  # Pula o título pois já foi adicionado
                                value = row[col]
                                if pd.notnull(value):
                                    formatted_value = self.format_value(value)
                                    if formatted_value.strip():  # Verifica se não está vazio
                                        p = doc.add_paragraph(formatted_value)
                                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        doc.add_paragraph()  # Linha em branco entre registros

                    # Salvar documento
                    doc.save(file_path)
                    logger.info(f"Arquivo DOCX salvo: {file_path}")
                else:
                    raise ValueError(f"Formato de arquivo não suportado: {file_format}")
                
                return {
                    "mensagem": "Scraping concluído com sucesso",
                    "total_itens": len(all_results),
                    "total_paginas": page,
                    "arquivo": f"{file_name}.{file_format}"
                }

            except Exception as e:
                logger.error(f"Erro ao salvar arquivo: {str(e)}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Erro ao salvar arquivo: {str(e)}"
                )

        except Exception as e:
            logger.error(f"Erro durante o scraping: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=str(e)
            )
        
        finally:
            if driver:
                driver.quit()
            
def format_value(self, value):
    """Formata valores para exibição no documento"""
    if pd.isna(value):
        return ''
    elif isinstance(value, (int, float)):
        return f"{value:,}"
    else:
        return str(value)
    
@app.post("/scraping")
async def scrape(request: ScrapingRequest):
    try:
        if not request.url:
            raise HTTPException(status_code=400, detail="URL é obrigatória")
        if not request.seletores or 'container' not in request.seletores:
            raise HTTPException(status_code=400, detail="Seletor 'container' é obrigatório")
        
        manager = ScrapingManager()
        return await manager.process_all_pages(request)  # Corrigido aqui
    
    except Exception as e:
        logger.error(f"Erro: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(EXPORTS_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, filename=filename)
    raise HTTPException(status_code=404, detail="Arquivo não encontrado")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)